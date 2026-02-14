"""Script to generate .docx templates with Jinja2 placeholders.

Run once: python scripts/create_templates.py
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")


def set_style(paragraph, font_size=11, bold=False, alignment=None):
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.name = "Times New Roman"
        run.bold = bold
    if alignment is not None:
        paragraph.alignment = alignment


def create_service_agreement():
    doc = Document()

    # Title
    p = doc.add_paragraph()
    run = p.add_run("ДОГОВОР ОКАЗАНИЯ УСЛУГ")
    set_style(p, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    run = p.add_run("№ {{ document_number }} от {{ contract_date }}")
    set_style(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # Preamble
    p = doc.add_paragraph()
    run = p.add_run(
        "{{ executor_name }}, ИНН {{ executor_inn }}, "
        "адрес: {{ executor_address }}, именуемый(-ая) в дальнейшем «Исполнитель», "
        "с одной стороны, и {{ client_name }}, ИНН {{ client_inn }}, "
        "адрес: {{ client_address }}, именуемый(-ая) в дальнейшем «Заказчик», "
        "с другой стороны, заключили настоящий договор о нижеследующем:"
    )
    set_style(p)

    # Section 1
    p = doc.add_paragraph()
    run = p.add_run("1. ПРЕДМЕТ ДОГОВОРА")
    set_style(p, bold=True)

    p = doc.add_paragraph()
    run = p.add_run(
        "1.1. Исполнитель обязуется по заданию Заказчика оказать следующие услуги: "
        "{{ service_description }}."
    )
    set_style(p)

    p = doc.add_paragraph()
    run = p.add_run(
        "1.2. Заказчик обязуется оплатить услуги в размере и порядке, "
        "предусмотренных настоящим договором."
    )
    set_style(p)

    # Section 2
    p = doc.add_paragraph()
    run = p.add_run("2. СТОИМОСТЬ УСЛУГ И ПОРЯДОК ОПЛАТЫ")
    set_style(p, bold=True)

    p = doc.add_paragraph()
    run = p.add_run(
        "2.1. Стоимость услуг по настоящему договору составляет "
        "{{ contract_amount }} ({{ contract_amount }}) руб."
    )
    set_style(p)

    p = doc.add_paragraph()
    run = p.add_run(
        "2.2. Оплата производится в течение 5 (пяти) рабочих дней "
        "с момента подписания акта выполненных работ."
    )
    set_style(p)

    # Section 3
    p = doc.add_paragraph()
    run = p.add_run("3. ПОРЯДОК СДАЧИ-ПРИЁМКИ УСЛУГ")
    set_style(p, bold=True)

    p = doc.add_paragraph()
    run = p.add_run(
        "3.1. По завершении оказания услуг Исполнитель направляет Заказчику "
        "акт выполненных работ."
    )
    set_style(p)

    p = doc.add_paragraph()
    run = p.add_run(
        "3.2. Заказчик обязан рассмотреть акт в течение 5 (пяти) рабочих дней "
        "и направить Исполнителю подписанный акт или мотивированный отказ."
    )
    set_style(p)

    # Section 4
    p = doc.add_paragraph()
    run = p.add_run("4. ОТВЕТСТВЕННОСТЬ СТОРОН")
    set_style(p, bold=True)

    p = doc.add_paragraph()
    run = p.add_run(
        "4.1. За неисполнение или ненадлежащее исполнение обязательств "
        "по настоящему договору стороны несут ответственность в соответствии "
        "с действующим законодательством Российской Федерации."
    )
    set_style(p)

    # Section 5
    p = doc.add_paragraph()
    run = p.add_run("5. СРОК ДЕЙСТВИЯ ДОГОВОРА")
    set_style(p, bold=True)

    p = doc.add_paragraph()
    run = p.add_run(
        "5.1. Настоящий договор вступает в силу с момента его подписания "
        "и действует до полного исполнения сторонами своих обязательств."
    )
    set_style(p)

    # Signatures
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("6. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
    set_style(p, bold=True)

    doc.add_paragraph()

    # Two-column signatures
    table = doc.add_table(rows=1, cols=2)
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)

    cell_left.text = (
        "ИСПОЛНИТЕЛЬ:\n\n"
        "{{ executor_name }}\n"
        "ИНН: {{ executor_inn }}\n"
        "Адрес: {{ executor_address }}\n\n\n"
        "___________________ / {{ executor_name }} /"
    )

    cell_right.text = (
        "ЗАКАЗЧИК:\n\n"
        "{{ client_name }}\n"
        "ИНН: {{ client_inn }}\n"
        "Адрес: {{ client_address }}\n\n\n"
        "___________________ / {{ client_name }} /"
    )

    path = os.path.join(TEMPLATES_DIR, "service_agreement.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_invoice():
    doc = Document()

    p = doc.add_paragraph()
    run = p.add_run("СЧЁТ НА ОПЛАТУ")
    set_style(p, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    run = p.add_run("№ {{ document_number }} от {{ invoice_date }}")
    set_style(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # Receiver info
    p = doc.add_paragraph()
    run = p.add_run("Получатель: {{ executor_name }}")
    set_style(p, bold=True)

    p = doc.add_paragraph()
    run = p.add_run("ИНН: {{ executor_inn }}")
    set_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Банк: {{ executor_bank }}")
    set_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Р/с: {{ executor_account }}")
    set_style(p)

    p = doc.add_paragraph()
    run = p.add_run("БИК: {{ executor_bik }}")
    set_style(p)

    doc.add_paragraph()

    # Payer info
    p = doc.add_paragraph()
    run = p.add_run("Плательщик: {{ client_name }}")
    set_style(p, bold=True)

    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"

    # Header
    cells = table.rows[0].cells
    cells[0].text = "№"
    cells[1].text = "Наименование"
    cells[2].text = "Сумма, руб."
    cells[3].text = "Итого, руб."

    # Data row
    cells = table.rows[1].cells
    cells[0].text = "1"
    cells[1].text = "{{ service_description }}"
    cells[2].text = "{{ amount }}"
    cells[3].text = "{{ amount }}"

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Итого к оплате: {{ amount }} руб.")
    set_style(p, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("___________________ / {{ executor_name }} /")
    set_style(p)

    path = os.path.join(TEMPLATES_DIR, "invoice.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_act_of_work():
    doc = Document()

    p = doc.add_paragraph()
    run = p.add_run("АКТ ВЫПОЛНЕННЫХ РАБОТ")
    set_style(p, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    run = p.add_run("№ {{ document_number }} от {{ act_date }}")
    set_style(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(
        "к Договору № {{ contract_number }} от {{ contract_date }}"
    )
    set_style(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # Preamble
    p = doc.add_paragraph()
    run = p.add_run(
        "{{ executor_name }}, именуемый(-ая) в дальнейшем «Исполнитель», "
        "с одной стороны, и {{ client_name }}, именуемый(-ая) в дальнейшем "
        "«Заказчик», с другой стороны, составили настоящий акт о нижеследующем:"
    )
    set_style(p)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(
        "1. Исполнитель выполнил, а Заказчик принял следующие работы/услуги:"
    )
    set_style(p)

    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"

    cells = table.rows[0].cells
    cells[0].text = "№"
    cells[1].text = "Наименование работ/услуг"
    cells[2].text = "Сумма, руб."

    cells = table.rows[1].cells
    cells[0].text = "1"
    cells[1].text = "{{ work_description }}"
    cells[2].text = "{{ amount }}"

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Итого: {{ amount }} руб.")
    set_style(p, bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(
        "2. Работы выполнены в полном объёме, в установленные сроки и "
        "с надлежащим качеством. Заказчик претензий по объёму, качеству "
        "и срокам выполнения работ не имеет."
    )
    set_style(p)

    doc.add_paragraph()

    # Signatures
    table = doc.add_table(rows=1, cols=2)
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)

    cell_left.text = (
        "ИСПОЛНИТЕЛЬ:\n\n\n"
        "___________________ / {{ executor_name }} /"
    )

    cell_right.text = (
        "ЗАКАЗЧИК:\n\n\n"
        "___________________ / {{ client_name }} /"
    )

    path = os.path.join(TEMPLATES_DIR, "act_of_work.docx")
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    create_service_agreement()
    create_invoice()
    create_act_of_work()
    print("All templates created successfully!")
