"""Parse company requisites from .docx and .pdf files."""

import fitz  # PyMuPDF
from docx import Document

REQUISITE_PROMPT = (
    "Ты — эксперт по анализу карточек предприятий и реквизитов организаций.\n"
    "Тебе дан текст документа с реквизитами. Извлеки следующие поля:\n\n"
    "- company_name: Полное наименование организации\n"
    "- legal_address: Юридический адрес\n"
    "- phone_email: Телефон / электронная почта / сайт (всё что есть)\n"
    "- ogrn: ОГРН\n"
    "- inn: ИНН\n"
    "- kpp: КПП\n"
    "- bank_account: Расчётный счёт\n"
    "- corr_account: Корреспондентский счёт\n"
    "- bik: БИК\n"
    "- bank_name: Название банка\n"
    "- bank_inn: ИНН банка\n"
    "- bank_address: Юридический адрес банка\n"
    "- director: ФИО генерального директора (только ФИО, без должности)\n\n"
    "Верни СТРОГО JSON без markdown и без ```json, только найденные поля:\n"
    '{"company_name": "...", "inn": "...", ...}\n'
    "Если поле не найдено в тексте — не включай его.\n"
)

# Maps AI-returned keys to possible template field keys.
# Each AI key maps to a list of candidate field keys, tried in order.
REQUISITE_TO_FIELD_MAP: dict[str, list[str]] = {
    "company_name": [
        "client_name", "executor_name",
        "customer_company_name", "executor_full_name",
    ],
    "legal_address": [
        "client_address", "executor_address",
        "customer_address",
    ],
    "phone_email": [
        "client_phone", "executor_phone",
        "customer_phone",
    ],
    "ogrn": [
        "client_ogrn", "executor_ogrn",
        "customer_ogrn", "executor_ogrnip",
    ],
    "inn": [
        "client_inn", "executor_inn",
        "customer_inn",
    ],
    "kpp": [
        "client_kpp", "executor_kpp",
        "customer_kpp",
    ],
    "bank_account": [
        "client_account", "executor_account",
        "customer_bank_rs", "executor_account_number",
    ],
    "corr_account": [
        "client_corr_account", "executor_corr_account",
        "customer_bank_ks", "executor_correspondent_account",
    ],
    "bik": [
        "client_bik", "executor_bik",
        "customer_bank_bik", "executor_bik",
    ],
    "bank_name": [
        "client_bank", "executor_bank",
        "customer_bank_name", "executor_bank_name",
    ],
    "bank_inn": [
        "executor_bank_inn", "customer_bank_inn",
    ],
    "bank_address": [
        "executor_bank_address", "customer_bank_address",
    ],
    "director": [
        "client_director", "executor_director",
        "customer_director_full_name",
    ],
}


def extract_text_from_pdf(path: str) -> str:
    """Extract text from all pages of a PDF using PyMuPDF."""
    doc = fitz.open(path)
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n".join(pages)


def extract_text_from_docx(path: str) -> str:
    """Extract text from .docx — paragraphs and tables as 'label: value' pairs."""
    doc = Document(path)
    lines = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Format as "label: value" for 2-column tables
            if len(cells) == 2 and cells[0] and cells[1]:
                lines.append(f"{cells[0]}: {cells[1]}")
            else:
                line = " | ".join(c for c in cells if c)
                if line:
                    lines.append(line)

    return "\n".join(lines)


def detect_side(fields: list[dict], current_index: int) -> str:
    """Determine if we're filling client or executor fields based on current field group.

    Returns a prefix that matches the field key prefixes: 'client', 'customer', or 'executor'.
    """
    if current_index < len(fields):
        group = fields[current_index].get("group", "").lower()
        if any(kw in group for kw in ("исполнитель", "получатель")):
            return "executor"
        # Check field key prefix to determine correct side prefix
        key = fields[current_index].get("key", "")
        if key.startswith("customer_"):
            return "customer"
    return "client"


def map_requisites_to_fields(
    requisites: dict,
    fields: list[dict],
    side: str,
) -> dict[str, str]:
    """Map parsed requisites to template field keys for the given side.

    Args:
        requisites: AI-parsed dict like {"company_name": "ООО ...", "inn": "123..."}
        fields: Template field definitions (each has "key")
        side: "client" or "executor"

    Returns:
        Dict mapping field_key -> value for fields that matched.
    """
    field_keys = {f["key"] for f in fields}
    result = {}

    for req_key, value in requisites.items():
        if not value:
            continue
        candidates = REQUISITE_TO_FIELD_MAP.get(req_key, [])
        for candidate in candidates:
            if candidate in field_keys and candidate.startswith(side):
                result[candidate] = str(value).strip()
                break

    return result


# Display labels for showing saved requisites to the user
REQUISITE_LABELS: dict[str, str] = {
    "company_name": "Наименование",
    "legal_address": "Юридический адрес",
    "phone_email": "Телефон / почта",
    "ogrn": "ОГРН",
    "inn": "ИНН",
    "kpp": "КПП",
    "bank_account": "Расчётный счёт",
    "corr_account": "Корр. счёт",
    "bik": "БИК",
    "bank_name": "Банк",
    "bank_inn": "ИНН банка",
    "bank_address": "Адрес банка",
    "director": "Генеральный директор",
}


def format_requisites_summary(requisites: dict) -> str:
    """Format saved requisites for display."""
    lines = []
    for key, label in REQUISITE_LABELS.items():
        val = requisites.get(key)
        if val:
            lines.append(f"│ {label}: {val}")
    return "\n".join(lines)
